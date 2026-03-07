from pathlib import Path
from typing import List
from collections import defaultdict
import docx
from .BaseParser import BaseParser

class DocxParser(BaseParser):
    """
    DOCX 文件解析器
    """
    type = 'docx'
    
    def _extract_content(self) -> str:
        """
        使用 python-docx 提取 DOCX 文件內容
        """
        try:
            doc = docx.Document(self.file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Error reading DOCX file {self.file_path}: {e}")
            return ""

    def _extract_metadata(self) -> defaultdict:
        metadata = defaultdict(str)
        try:
            doc = docx.Document(self.file_path)
            core_props = doc.core_properties
            metadata['title'] = core_props.title
            metadata['author'] = core_props.author
            metadata['subject'] = core_props.subject
            metadata['keywords'] = core_props.keywords
            metadata['created'] = str(core_props.created)
            metadata['modified'] = str(core_props.modified)
        except Exception:
            pass
        return metadata

    def _check_format(self) -> bool:
        f_path: Path = Path(self.file_path)
        return f_path.exists() and f_path.suffix.lower() == '.docx'
