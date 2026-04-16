import os
import docx
from .TextChunkProcessor import TextChunkProcessor


class DocParser(TextChunkProcessor):
    """
    Word 文件解析器（支持 .doc 和 .docx）
    
    解析方案优先级：
    - .docx: python-docx (轻量、跨平台)
    - .doc:   doc2txt (轻量) -> win32com (需Office) -> textract (需系统依赖)
    
    特性：
    - 提取表格内容并转换为 Markdown 格式
    """
    type = 'doc'

    def _extract_content(self) -> str:
        ext = os.path.splitext(self.file_path)[1].lower()

        if ext == '.docx':
            return self._extract_with_docx()

        if ext == '.doc':
            # 方案1: doc2txt（轻量，只需 pip install）
            content = self._extract_with_doc2txt()
            if content:
                return content

            # 方案2: win32com（仅Windows，需要Office）
            content = self._extract_with_win32com()
            if content:
                return content

            # 方案3: textract（需系统依赖：antiword 或 LibreOffice）
            content = self._extract_with_textract()
            if content:
                return content

        return ""

    def _extract_with_docx(self) -> str:
        """使用 python-docx 提取文本和表格（仅支持 .docx）"""
        try:
            doc = docx.Document(self.file_path)
            content_parts = []

            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text.strip())

            # 提取表格
            tables = self._extract_tables(doc)
            content_parts.extend(tables)

            return '\n'.join(content_parts)
        except Exception:
            return ""

    def _extract_tables(self, doc) -> list:
        """提取 Word 文档中的所有表格，转换为 Markdown 格式"""
        tables = []
        
        for table in doc.tables:
            if not table.rows:
                continue

            # 收集表格数据
            rows_data = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_data.append(cells)

            if not rows_data:
                continue

            # 转换为 Markdown 表格
            table_lines = []
            
            # 表头
            table_lines.append('| ' + ' | '.join(rows_data[0]) + ' |')
            
            # 分隔行
            col_count = len(rows_data[0])
            table_lines.append('| ' + ' | '.join(['---'] * col_count) + ' |')
            
            # 数据行
            for row in rows_data[1:]:
                table_lines.append('| ' + ' | '.join(row) + ' |')

            tables.append('\n'.join(table_lines))

        return tables

    def _extract_with_doc2txt(self) -> str:
        """使用 doc2txt 提取文本（轻量，纯 Python）"""
        try:
            import doc2txt
            text = doc2txt.extract_text(self.file_path)
            return text.strip() if text else ""
        except ImportError:
            return ""
        except Exception:
            return ""

    def _extract_with_textract(self) -> str:
        """使用 textract 提取文本（支持 .doc 和 .docx，需系统依赖）"""
        try:
            import textract
            text = textract.process(self.file_path, encoding='utf-8').decode('utf-8')
            return text.strip()
        except ImportError:
            return ""
        except Exception:
            return ""

    def _extract_with_win32com(self) -> str:
        """使用 win32com 提取文本（仅Windows，需要Office）"""
        try:
            import win32com.client
        except ImportError:
            return ""

        try:
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            try:
                doc = word.Documents.Open(self.file_path)
                try:
                    full_text = [para.Range.Text.strip() for para in doc.Paragraphs if para.Range.Text.strip()]
                    return '\n'.join(full_text)
                finally:
                    doc.Close(False)
            finally:
                word.Quit()
        except Exception:
            return ""
